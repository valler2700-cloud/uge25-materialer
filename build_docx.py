# -*- coding: utf-8 -*-
"""
Markdown -> DOCX konverter for Uge25-materialerne.
Understøtter: overskrifter, fed/kursiv, klikbare links (også bare URL'er),
punkt- og nummerlister (med ombrudte linjer), tabeller, citat-/infobokse,
vandrette streger, billeder (![cap](sti)) og billed-galleri ([[GALLERY ...]]).
Billeder nedskaleres automatisk med Pillow.
"""
import re, io, sys, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))

# ---------- inline (fed / kursiv / links / bare URL) ----------
INLINE = re.compile(
    r'(\[[^\]]+\]\([^)]+\)'      # [text](url)
    r'|\*\*[^*]+\*\*'            # **bold**
    r'|\*[^*]+\*'               # *italic*
    r'|https?://[^\s)]+)'        # bare url
)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t); link.append(run); paragraph._p.append(link)

def add_runs(paragraph, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('[') and '](' in tok:
            lt = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            add_hyperlink(paragraph, lt.group(2), lt.group(1))
        elif tok.startswith('**'):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*'):
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith('http'):
            add_hyperlink(paragraph, tok, tok)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

# ---------- helpers ----------
def shade(p, fill='EEF1F6'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.10)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'BBBBBB')
    bdr.append(b); pPr.append(bdr)

def img_stream(path, width_in):
    im = Image.open(path)
    if im.mode in ('RGBA', 'P', 'LA'):
        im = im.convert('RGB')
    target = int(width_in * 240)
    if im.width > target:
        im = im.resize((target, int(im.height * target / im.width)))
    bio = io.BytesIO(); im.save(bio, format='JPEG', quality=85); bio.seek(0)
    return bio

def add_image(doc, path, cap, width_in=5.3):
    full = os.path.join(BASE, path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_stream(full, width_in), width=Inches(width_in))
    if cap:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(cap); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def add_gallery(doc, spec):
    # spec: "w=1.4 | path :: cap | path :: cap | ..."
    m = re.match(r'\s*w=([\d.]+)\s*\|(.*)', spec, re.S)
    width = float(m.group(1)); items = [s.strip() for s in m.group(2).split('|') if s.strip()]
    parsed = []
    for it in items:
        path, _, cap = it.partition('::')
        parsed.append((path.strip(), cap.strip()))
    table = doc.add_table(rows=1, cols=len(parsed))
    table.alignment = 1
    for cell, (path, cap) in zip(table.rows[0].cells, parsed):
        pic_p = cell.paragraphs[0]; pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(img_stream(os.path.join(BASE, path), width), width=Inches(width))
        capp = cell.add_paragraph(); capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = capp.add_run(cap); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    header, body = cells[0], cells[2:]
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]; cell.paragraphs[0].text = ''
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs: run.bold = True
    for row in body:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            if i < len(rc):
                rc[i].paragraphs[0].text = ''
                add_runs(rc[i].paragraphs[0], val)

def is_block(line):
    s = line.strip()
    return (s == '' or s == '---' or line.startswith('#') or line.startswith('>')
            or line.lstrip().startswith('|') or s.startswith('![') or s.startswith('[[GALLERY')
            or re.match(r'^\s*[-*]\s', line) or re.match(r'^\s*\d+\.\s', line))

# ---------- main convert ----------
def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == '':
            i += 1; continue
        if s == '---':
            add_hr(doc); i += 1; continue
        if line.startswith('#'):
            m = re.match(r'(#+)\s+(.*)', line)
            lvl = min(len(m.group(1)), 4)
            h = doc.add_heading('', level=lvl)
            add_runs(h, m.group(2).strip()); i += 1; continue
        if s.startswith('[[GALLERY'):
            spec = s[len('[[GALLERY'):].rstrip(']').strip()
            add_gallery(doc, spec); i += 1; continue
        im = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', s)
        if im:
            add_image(doc, im.group(2).strip(), im.group(1).strip()); i += 1; continue
        if line.startswith('>'):
            block = []
            while i < n and lines[i].startswith('>'):
                block.append(lines[i].lstrip('>').strip()); i += 1
            # join into paragraphs split by empty quote lines
            para = []
            for q in block + ['']:
                if q == '':
                    if para:
                        p = doc.add_paragraph(); shade(p)
                        add_runs(p, ' '.join(para)); para = []
                else:
                    para.append(q.lstrip('_').rstrip('_'))
            continue
        if line.lstrip().startswith('|'):
            block = []
            while i < n and lines[i].lstrip().startswith('|'):
                block.append(lines[i]); i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue
        mlist = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if mlist:
            style = 'List Number' if mlist.group(2)[0].isdigit() else 'List Bullet'
            while i < n:
                ml = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', lines[i])
                if not ml:
                    break
                txt = ml.group(3)
                i += 1
                # gather wrapped continuation lines
                while i < n and lines[i].strip() != '' and not is_block(lines[i]):
                    txt += ' ' + lines[i].strip(); i += 1
                p = doc.add_paragraph(style=style)
                add_runs(p, txt)
            continue
        # plain paragraph (join wrapped lines)
        para = [s]; i += 1
        while i < n and lines[i].strip() != '' and not is_block(lines[i]):
            para.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_runs(p, ' '.join(para))
    doc.save(docx_path)
    print('OK ->', os.path.basename(docx_path))

import glob

def build_all():
    """Byg alle *.md i scriptets mappe til matchende .docx (springer README over)."""
    SKIP = {'README.md'}
    built = 0
    for md in sorted(glob.glob(os.path.join(BASE, '*.md'))):
        if os.path.basename(md) in SKIP:
            continue
        convert(md, md[:-3] + '.docx')
        built += 1
    print(f'Byggede {built} dokument(er).')

if __name__ == '__main__':
    pairs = sys.argv[1:]
    if pairs:
        for j in range(0, len(pairs), 2):
            convert(pairs[j], pairs[j + 1])
    else:
        build_all()
